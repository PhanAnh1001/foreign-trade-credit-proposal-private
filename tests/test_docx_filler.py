"""Tests for DOCX template filler."""
import os
import tempfile
import pytest
from pathlib import Path
from src.utils.docx_filler import fill_lc_template

TEMPLATE = Path("data/templates/docx/vietcombank/Application-for-LC-issuance.docx")


def _sample_data() -> dict:
    return {
        "applicant_name": "Viet Nam Technology Import-Export JSC",
        "applicant_address": "123 Nguyen Thi Minh Khai, District 1, Ho Chi Minh City",
        "beneficiary_name": "Shenzhen Advanced Electronics Co., Ltd.",
        "beneficiary_address": "No. 88 Technology Park Road, Nanshan, Shenzhen, China",
        "beneficiary_account_no": "6227 0036 8900 1234567",
        "beneficiary_bank_name": "Bank of China, Shenzhen Nanshan Branch",
        "beneficiary_bank_address": "No. 12 Chuangye Road, Nanshan District, Shenzhen",
        "beneficiary_bank_bic": "BKCHCNBJ510",
        "lc_type": "Irrevocable",
        "issuance_method": "SWIFT",
        "currency": "USD",
        "amount": "450000.00",
        "amount_in_words": "SAY US DOLLARS FOUR HUNDRED FIFTY THOUSAND ONLY",
        "amount_tolerance": "0",
        "expiry_date": "25/02/28",
        "expiry_place": "At the counter of the issuing bank (Vietcombank)",
        "latest_shipment_date": "25/01/31",
        "incoterms": "CIF",
        "incoterms_version": "2020",
        "named_port": "Ho Chi Minh City Port, Vietnam",
        "port_of_loading": "Shekou Port, Shenzhen, China",
        "port_of_discharge": "Cat Lai Port, Ho Chi Minh City, Vietnam",
        "partial_shipment": "Not allowed",
        "transhipment": "Not allowed",
        "draft_type": "Sight",
        "draft_days": None,
        "presentation_period": "21",
        "description_of_goods": (
            "Electronic Circuit Boards and Semiconductor Components (AEC-2024 Series), "
            "HS Code: 8542.31.10, 1,000 units @ USD 450.00/unit"
        ),
        "documents": {
            "commercial_invoice": "3 originals, signed by Seller",
            "bill_of_lading": (
                "Full set of 3/3 original clean shipped on board ocean bills of lading, "
                "made out to order, notify applicant, marked 'Freight Prepaid'"
            ),
            "packing_list": "1 original + 2 copies",
            "certificate_of_origin": "1 original (Form E or CCPIT certificate)",
            "insurance_certificate": (
                "1 original insurance certificate/policy covering all risks "
                "(Institute Cargo Clauses A), 110% of invoice value, in USD"
            ),
            "inspection_certificate": None,
            "other_documents": [],
        },
        "additional_conditions": (
            "Documents must be issued in English\n"
            "The amount utilized must be endorsed on the reverse of the original L/C."
        ),
        "issuing_bank_charges_for": "Applicant",
        "other_bank_charges_for": "Beneficiary",
        "contract_number": "VN-CN-2024-001",
        "contract_date": "01/11/2024",
        "vcb_branch": "Ho Chi Minh City Branch",
        "validation_warnings": [],
        "compliance_notes": [],
    }


@pytest.mark.skipif(not TEMPLATE.exists(), reason="LC template not found")
def test_fill_creates_output():
    data = _sample_data()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    try:
        result = fill_lc_template(data, str(TEMPLATE), output_path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 1000  # non-trivial file
    finally:
        os.unlink(output_path)


@pytest.mark.skipif(not TEMPLATE.exists(), reason="LC template not found")
def test_fill_contains_applicant_name():
    from docx import Document
    data = _sample_data()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    try:
        fill_lc_template(data, str(TEMPLATE), output_path)
        doc = Document(output_path)
        full_text = "\n".join(
            p.text for table in doc.tables
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        )
        assert "Viet Nam Technology Import-Export JSC" in full_text
    finally:
        os.unlink(output_path)


@pytest.mark.skipif(not TEMPLATE.exists(), reason="LC template not found")
def test_fill_contains_currency_and_amount():
    from docx import Document
    data = _sample_data()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    try:
        fill_lc_template(data, str(TEMPLATE), output_path)
        doc = Document(output_path)
        full_text = "\n".join(
            p.text for table in doc.tables
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        )
        assert "USD" in full_text
        assert "450000" in full_text
    finally:
        os.unlink(output_path)
