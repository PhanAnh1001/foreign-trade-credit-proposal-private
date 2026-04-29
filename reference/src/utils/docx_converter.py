import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    """Convert Markdown string to a DOCX file.

    Handles headers (#, ##, ###), bold, italic, tables, bullet lists,
    numbered lists, horizontal rules, and normal paragraphs.

    Returns the output file path.
    """
    doc = Document()

    # Set default font for Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1 (single #, not ##)
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1

        # H2 (## but not ###)
        elif line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1

        # H3
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1

        # Table: collect consecutive rows starting with '|'
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)

        # Bullet list (- or *)
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _parse_inline(para, line[2:].strip())
            i += 1

        # Numbered list (1. 2. etc.)
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            para = doc.add_paragraph(style='List Number')
            _parse_inline(para, text.strip())
            i += 1

        # Horizontal rule
        elif re.match(r'^-{3,}$', line.strip()) or re.match(r'^_{3,}$', line.strip()):
            doc.add_paragraph('_' * 50)
            i += 1

        # Empty line — skip (natural paragraph break)
        elif line.strip() == '':
            i += 1

        # Normal paragraph
        else:
            para = doc.add_paragraph()
            _parse_inline(para, line.strip())
            i += 1

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _parse_inline(para, text: str) -> None:
    """Parse inline markdown formatting (**bold**, *italic*) and add runs to para."""
    # Pattern matches **bold** or *italic* tokens
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _apply_inline_formatting(para) -> None:
    """Apply bold/italic inline formatting to an existing paragraph's text.

    Clears existing runs and re-parses the combined text for inline markers.
    """
    full_text = ''.join(r.text for r in para.runs)
    if not full_text:
        return
    para.clear()
    _parse_inline(para, full_text)


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    """Parse markdown table lines and add as a DOCX table with Table Grid style."""
    rows = []
    for line in table_lines:
        # Skip separator lines: |---|---| or |:---|:---:|
        if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            para = cell.paragraphs[0]
            _parse_inline(para, cell_text)
            # Bold the header row
            if row_idx == 0:
                for run in para.runs:
                    run.bold = True
