"""DOCX Template Renderer.

Two public functions:
  render_from_template() — fills VPBank form template with Output 1 (company info)
  render_analyst_memo()  — creates a new analyst memo DOCX with Output 2 + 3

Cell coordinate reference: docs/requirements/docx_template_design.md
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

from ..models.company import CompanyInfo
from ..models.financial import FinancialData
from ..utils.logger import get_logger
from ..config import FORM_TEMPLATE_DOCX

logger = get_logger("docx_template")

_TEMPLATE_PATH = str(FORM_TEMPLATE_DOCX)


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def render_from_template(
    output_path: str,
    company_info: CompanyInfo | None = None,
    financial_data: FinancialData | None = None,
    template_path: str = _TEMPLATE_PATH,
) -> str:
    """Fill VPBank DOCX template with real data and save to output_path.

    Output DOCX preserves template structure exactly (35 tables, 19 sections,
    13 images, green borders). No new sections are appended — Phụ lục A/B
    analysis text is only written to the markdown output, not to the DOCX.

    Args:
        output_path:    Destination DOCX file path.
        company_info:   Extracted company info (CompanyInfo model).
        financial_data: Extracted financial data (FinancialData model).
        template_path:  Source DOCX template (default: giay-de-nghi-vay-von.docx).

    Returns:
        output_path on success.
    """
    logger.info(f"Loading template: {template_path}")
    doc = Document(template_path)

    if company_info:
        _fill_company_info(doc, company_info)
        _fill_shareholders(doc, company_info)
        _fill_phu_luc_1(doc, company_info)
        _fill_board_and_management(doc, company_info)
        _fill_business_ops(doc, company_info)

    if financial_data and financial_data.statements:
        _fill_vot_thuc_gop(doc, financial_data)
        _fill_financial_history(doc, financial_data)
        _fill_income_statement(doc, financial_data)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved → {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Section fillers
# ─────────────────────────────────────────────────────────────────────────────

def _fill_company_info(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[1] (1.1 legal entity) and Table[2] row 14 (representative).

    Table[1] layout: 13 rows × 2 cols — col 0 = label, col 1 = value cell.
    """
    t1 = doc.tables[1]

    _set_cell(t1.rows[2].cells[1], info.company_name)        # Tên Khách hàng
    # Tax code + ngày cấp + cơ quan cấp → right cell (cells[1])
    if info.tax_code:
        reg_value = info.tax_code
        if info.established_date:
            reg_value += f" Ngày cấp: {info.established_date}"
        if info.registration_authority:
            reg_value += f" Cơ quan cấp: {info.registration_authority}"
        _set_cell(t1.rows[3].cells[1], reg_value)
    _set_cell(t1.rows[4].cells[1], info.address)             # Địa chỉ trụ sở
    _set_cell(t1.rows[5].cells[1], info.address)             # Địa chỉ giao dịch (same)
    _set_cell(t1.rows[6].cells[1], info.phone)               # Điện thoại
    _set_cell(t1.rows[7].cells[1], info.main_business)       # Ngành nghề KD chính
    _set_cell(t1.rows[9].cells[1], info.charter_capital)     # Vốn điều lệ

    # Table[2]: row 14 col 1 = Người đại diện pháp luật
    if info.legal_representative:
        t2 = doc.tables[2]
        _set_cell(t2.rows[14].cells[1], info.legal_representative)

    logger.debug("Company info filled (Table[1] + Table[2] r14)")


def _fill_shareholders(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[3] rows 10–12 — Section 1.3 cổ đông/thành viên góp vốn.

    Template has 3 pre-existing empty data rows (r10, r11, r12).
    Columns (9-col table, some merged): [0]=STT [1+2]=Họ tên [3+4]=Mối QH
    [5]=Tỷ lệ% [6+7]=Kinh nghiệm [8]=Dư nợ VPB.
    """
    if not info.shareholders:
        return

    t3 = doc.tables[3]
    for i, sh in enumerate(info.shareholders[:3]):
        row = t3.rows[10 + i]
        _set_cell(row.cells[0], str(i + 1))           # STT
        _set_cell(row.cells[1], sh.name)               # Họ và tên (c1+c2 merged)
        if sh.percentage is not None:
            _set_cell(row.cells[5], f"{sh.percentage:.1f}%")  # Tỷ lệ góp vốn
        # cells[3] = "Mối quan hệ" — left blank (no relationship data available)

    logger.debug(f"Shareholders filled: {min(len(info.shareholders), 3)} rows")


def _fill_phu_luc_1(doc: Document, info: CompanyInfo) -> None:
    """Fill PHỤ LỤC 1 with the primary (largest) shareholder — individual section.

    PHỤ LỤC 1 is a FORM (label | value rows), not a data table.
    Template structure:
      Table[14] r0–r13 : Section Doanh nghiệp (skipped — MST shareholders are individuals)
      Table[14] r14    : Section header "Thành viên góp vốn chính (nếu là cá nhân)"
      Table[14] r15 c1 : Mối quan hệ với khách hàng
      Table[14] r16 c1 : Họ và tên
      Table[15] r5  c1 : Tỷ lệ góp vốn

    Fills only shareholders[0] (the primary/largest shareholder).
    """
    if not info.shareholders:
        return

    sh = info.shareholders[0]

    try:
        t14 = doc.tables[14]
        t15 = doc.tables[15]
    except IndexError:
        logger.warning("Table[14]/[15] not found — skipping PHỤ LỤC 1 fill")
        return

    # Table[14] r15 col1: Mối quan hệ với khách hàng → fixed "Cổ đông chính"
    if len(t14.rows) > 15:
        _set_cell(t14.rows[15].cells[1], "Cổ đông chính")

    # Table[14] r16 col1: Họ và tên
    if len(t14.rows) > 16:
        _set_cell(t14.rows[16].cells[1], sh.name)

    # Table[15] r5 col1: Tỷ lệ góp vốn
    if len(t15.rows) > 5 and sh.percentage is not None:
        _set_cell(t15.rows[5].cells[1], f"{sh.percentage:.1f}%")

    logger.debug(f"PHỤ LỤC 1 (Table[14]/[15]) filled for primary shareholder: {sh.name}")


def _fill_vot_thuc_gop(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[1] row 10 col 1 — Vốn thực góp đến ngày...

    human_mapping note #1: Vốn thực góp = Vốn chủ sở hữu (equity, mã 400 CĐKT).
    Uses the latest year available in financial_data.
    """
    years = sorted(financial_data.statements.keys())
    if not years:
        return
    s = financial_data.statements[years[-1]]
    if s.equity is None:
        return

    t1 = doc.tables[1]
    if len(t1.rows) > 10:
        _set_cell(t1.rows[10].cells[1], _fmt(s.equity))
        logger.debug(f"Vốn thực góp (Table[1] r10) filled: {_fmt(s.equity)} (equity {years[-1]})")


def _fill_board_and_management(doc: Document, info: CompanyInfo) -> None:
    """No-op: HĐQT/BGĐ/BKS are not part of the VPBank loan application form.

    Form section 1.2 only captures the legal representative (filled in
    _fill_company_info via Table[2] r14). Board/management data belongs
    in the analyst memo, not this customer-facing form.
    """
    logger.debug("Board/management injection skipped — not applicable to loan form")


def _fill_business_ops(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[8] row 1 col 0 — Lĩnh vực kinh doanh chính."""
    if not info.main_business:
        return
    _set_cell(doc.tables[8].rows[1].cells[0], info.main_business)
    logger.debug("Business ops (Table[8]) filled")


def _fill_financial_history(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[29] — PHỤ LỤC 6: historical revenue / costs / net profit.

    Template: [0]=TT [1]=Chỉ tiêu [2]=Năm N-1 [3]=Năm kế hoạch [4]=Ghi chú
    Renderer renames col headers to actual years and fills with real data.
    Uses second-to-last year for col 2 and last year for col 3.
    """
    t29 = doc.tables[29]
    years = sorted(financial_data.statements.keys())
    if not years:
        return

    year_n1 = years[-2] if len(years) >= 2 else None
    year_n  = years[-1]

    # Rename column headers to actual years
    if year_n1:
        _set_cell(t29.rows[0].cells[2], str(year_n1))
    _set_cell(t29.rows[0].cells[3], str(year_n))

    for col_idx, year in [(2, year_n1), (3, year_n)]:
        if year is None:
            continue
        s = financial_data.statements.get(year)
        if not s:
            continue

        # Row 1: Doanh thu
        _set_cell(t29.rows[1].cells[col_idx], _fmt(s.net_revenue))

        # Row 2: Tổng chi phí = COGS + bán hàng + quản lý
        total_costs = sum(
            x for x in [s.cost_of_goods_sold, s.selling_expenses, s.admin_expenses]
            if x is not None
        )
        _set_cell(t29.rows[2].cells[col_idx], _fmt(total_costs or None))

        # Row 3: Lợi nhuận sau thuế
        _set_cell(t29.rows[3].cells[col_idx], _fmt(s.net_profit))

        # Row 4: Tổng nhu cầu vốn lưu động = Tài sản ngắn hạn - Nợ ngắn hạn (working capital)
        if s.current_assets is not None and s.current_liabilities is not None:
            working_capital = s.current_assets - s.current_liabilities
            _set_cell(t29.rows[4].cells[col_idx], _fmt(working_capital))

        # Row 5: Nguồn vốn tự có = Vốn chủ sở hữu (human_mapping note #1)
        _set_cell(t29.rows[5].cells[col_idx], _fmt(s.equity))

        # Row 6: Nhu cầu vốn vay TCTD khác = Tổng nợ phải trả (proxy, human_mapping note #2)
        _set_cell(t29.rows[6].cells[col_idx], _fmt(s.total_liabilities))

    logger.debug(f"Financial history (Table[29]) filled — years: {year_n1}, {year_n} (rows 1–6)")


def _fill_income_statement(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[32] — PHỤ LỤC 6: income statement detail for the latest year.

    Table[32] layout: [0]=STT [1]=Chi tiết [2]=12-month value [3]=Ghi chú
    Rows: 0=header 1=revenue 2=COGS 3=gross_profit 4=other_income
          5=finance_costs 6=tax 7=other_costs 8=net_profit
    """
    years = sorted(financial_data.statements.keys())
    if not years:
        return
    year = years[-1]
    s = financial_data.statements[year]

    t32 = doc.tables[32]
    _set_cell(t32.rows[0].cells[2], f"Năm {year}")          # Update header

    _set_cell(t32.rows[1].cells[2], _fmt(s.net_revenue))
    _set_cell(t32.rows[2].cells[2], _fmt(s.cost_of_goods_sold))
    _set_cell(t32.rows[3].cells[2], _fmt(s.gross_profit))
    # rows 4-6: other income / finance costs / tax — no data, leave blank
    other_costs = (s.admin_expenses or 0) + (s.selling_expenses or 0)
    _set_cell(t32.rows[7].cells[2], _fmt(other_costs or None))  # Chi phí khác
    _set_cell(t32.rows[8].cells[2], _fmt(s.net_profit))

    logger.debug(f"Income statement (Table[32]) filled for year {year}")


def render_analyst_memo(
    output_path: str,
    company_name: str,
    section2: str | None = None,
    section3: str | None = None,
    date: str | None = None,
) -> str:
    """Create a credit analyst memo DOCX (tờ trình thẩm định tín dụng nội bộ).

    Creates a fresh Document (not from template) with:
      - Cover block: company name, date, classification
      - Output 2: sector analysis (from section2 markdown)
      - Output 3: financial analysis (from section3 markdown)

    Args:
        output_path:  Destination DOCX path.
        company_name: Company name for the header.
        section2:     Sector analysis markdown (Output 2 from subgraph2).
        section3:     Financial analysis markdown (Output 3 from subgraph3).
        date:         Report date. Defaults to today DD/MM/YYYY.

    Returns:
        output_path on success.
    """
    if date is None:
        date = datetime.now().strftime("%d/%m/%Y")

    doc = Document()

    # ── Cover block ─────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("TỜ TRÌNH THẨM ĐỊNH TÍN DỤNG NỘI BỘ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # blank spacer

    doc.add_paragraph(f"Khách hàng: {company_name}")
    doc.add_paragraph(f"Ngày lập: {date}")
    doc.add_paragraph("Phân loại: Nội bộ – Bảo mật")

    doc.add_paragraph()  # blank spacer

    # ── Output 2: Sector analysis ────────────────────────────────────────────
    if section2 and section2.strip():
        _append_markdown(doc, section2)
        logger.debug("Sector section appended to analyst memo")

    # ── Output 3: Financial analysis ─────────────────────────────────────────
    if section3 and section3.strip():
        _add_page_break(doc)
        _append_markdown(doc, section3)
        logger.debug("Financial section appended to analyst memo")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Analyst memo DOCX saved → {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Low-level cell helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell(cell, text: str | None) -> None:
    """Clear all runs in a cell and set new text, preserving paragraph style."""
    if text is None:
        return
    text = str(text).strip()
    if not text:
        return
    # Clear existing run content across all paragraphs
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    # Write into first paragraph
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
    else:
        cell.add_paragraph(text)


def _append_cell(cell, text: str) -> None:
    """Append text to the last paragraph of a cell without clearing existing content."""
    if not text:
        return
    if cell.paragraphs:
        cell.paragraphs[-1].add_run(str(text))
    else:
        cell.add_paragraph(str(text))


# ─────────────────────────────────────────────────────────────────────────────
# Formatting helpers
# ─────────────────────────────────────────────────────────────────────────────

def _fmt(value: float | None, unit: str = "triệu đồng") -> str:
    """Format a financial value: None → '', ≥1000 → 'X,XXX.X tỷ đồng'."""
    if value is None or value == 0:
        return ""
    if abs(value) >= 1000:
        return f"{value / 1000:,.1f} tỷ đồng"
    return f"{value:,.1f} {unit}"


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → DOCX append
# ─────────────────────────────────────────────────────────────────────────────

def _add_heading_safe(doc: Document, text: str, level: int) -> None:
    """Add a heading paragraph.

    python-docx's add_heading() uses BabelFish.ui2internal() which lowercases
    'Heading 1' → 'heading 1'. The template stores the style as 'Heading 1'
    (capital H), so lookup by name fails. Workaround: look up by style_id
    ('Heading1'), suppress the deprecation warning, then set text.
    """
    import warnings
    style_id = f"Heading{level}"
    para = doc.add_paragraph()
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            para.style = doc.styles[style_id]
        para.add_run(text)
    except KeyError:
        # Fallback: bold paragraph if heading style not defined in template
        run = para.add_run(text)
        run.bold = True


def _append_markdown(doc: Document, md_text: str) -> None:
    """Parse markdown and append as DOCX elements (headings, tables, lists, text)."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            _add_heading_safe(doc, line[2:].strip(), level=1)
            i += 1

        elif line.startswith("## ") and not line.startswith("### "):
            _add_heading_safe(doc, line[3:].strip(), level=2)
            i += 1

        elif line.startswith("### "):
            _add_heading_safe(doc, line[4:].strip(), level=3)
            i += 1

        elif line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _md_table_to_docx(doc, tbl_lines)

        elif line.startswith("- ") or line.startswith("* "):
            # Template has no "List Bullet" style — use "List Paragraph" with bullet prefix
            para = doc.add_paragraph(style="List Paragraph")
            para.add_run("• " + line[2:].strip())
            i += 1

        elif re.match(r"^\d+\. ", line):
            # Template has no "List Number" style — keep number prefix in "List Paragraph"
            para = doc.add_paragraph(style="List Paragraph")
            para.add_run(line.strip())
            i += 1

        elif re.match(r"^-{3,}$", line.strip()):
            i += 1  # skip horizontal rules

        elif line.strip() == "":
            i += 1

        else:
            _add_paragraph_with_inline(doc, line.strip())
            i += 1


def _add_paragraph_with_inline(doc: Document, text: str) -> None:
    """Add a normal paragraph, parsing **bold** and *italic* inline markers."""
    para = doc.add_paragraph()
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    for part in re.split(pattern, text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def _md_table_to_docx(doc: Document, table_lines: list[str]) -> None:
    """Convert markdown table lines into a DOCX Table Grid table."""
    rows = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue  # skip separator rows
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    # Template only has "Table Normal" — look up by style_id "TableGrid" if available,
    # otherwise fall back to "Table Normal" (no visible borders, content still readable)
    import warnings
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            tbl.style = doc.styles["TableGrid"]
    except KeyError:
        try:
            tbl.style = doc.styles["Table Normal"]
        except KeyError:
            pass

    for ri, row_data in enumerate(rows):
        for ci in range(num_cols):
            text = row_data[ci] if ci < len(row_data) else ""
            cell = tbl.cell(ri, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            if ri == 0:
                run.bold = True
