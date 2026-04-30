"""Fill the Vietcombank LC application DOCX template with extracted LC data."""
from __future__ import annotations
import re
import shutil
from pathlib import Path
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


# ─── Low-level paragraph helpers ────────────────────────────────────────────

def _full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _set_para_text(para, new_text: str) -> None:
    """Replace paragraph text preserving formatting of the first non-empty run."""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _replace_in_para(para, old: str, new: str) -> bool:
    """Replace first occurrence of `old` in paragraph (handles multi-run text)."""
    full = _full_text(para)
    if old not in full:
        return False
    _set_para_text(para, full.replace(old, new, 1))
    return True


def _replace_in_cell(cell, old: str, new: str) -> bool:
    """Replace first occurrence of `old` anywhere in a table cell."""
    for para in cell.paragraphs:
        if _replace_in_para(para, old, new):
            return True
    return False


def _replace_in_doc(doc: Document, old: str, new: str) -> bool:
    """Replace `old` anywhere in document paragraphs and all table cells."""
    changed = False
    for para in doc.paragraphs:
        if _replace_in_para(para, old, new):
            changed = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_in_para(para, old, new):
                        changed = True
    return changed


# The template uses Wingdings font with  (open square) for unchecked boxes.
# We also handle plain Unicode □ (U+25A1) as a fallback.
_UNCHECKED = frozenset(['', '□', '◻'])
_CHECKED = '■'  # U+25A0 Black Square — universally supported


def _check_run0_in_para(para) -> bool:
    """Tick the Wingdings checkbox stored in Run 0 of a paragraph."""
    if para.runs and para.runs[0].text in _UNCHECKED:
        para.runs[0].text = _CHECKED
        return True
    return False


def _select_checkbox(para, option_text: str) -> bool:
    """Mark the checkbox immediately before `option_text` as selected.

    The template stores checkboxes as Wingdings \\uf06f in separate runs:
      Run i  = '\\uf06f'  (unchecked box)
      Run i+1 = ' '
      Run i+2 = 'Irrevocable'   ← option_text (must be exact single-run match)

    We find the run whose stripped text matches option_text, then look
    back up to 3 runs for the checkbox character and replace it.
    """
    runs = para.runs
    for i, run in enumerate(runs):
        if run.text.strip() == option_text:
            for j in range(max(0, i - 3), i):
                if any(c in runs[j].text for c in _UNCHECKED):
                    for c in _UNCHECKED:
                        if c in runs[j].text:
                            runs[j].text = runs[j].text.replace(c, _CHECKED, 1)
                            return True
    return False


def _select_checkbox_in_cell(cell, option_text: str) -> bool:
    for para in cell.paragraphs:
        if _select_checkbox(para, option_text):
            return True
    return False


def _select_nth_checkbox_in_cell(cell, n: int) -> bool:
    """Select the n-th (0-indexed) unchecked Wingdings checkbox in a cell."""
    count = 0
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text in _UNCHECKED:
                if count == n:
                    run.text = _CHECKED
                    return True
                count += 1
    return False


# ─── Field-specific fill functions ──────────────────────────────────────────

def _fill_header(doc: Document, data: dict) -> None:
    """Table 0: Bank branch, applicant name, CIF."""
    t0 = doc.tables[0]
    _replace_in_cell(t0.rows[0].cells[0], "………………………", data.get("vcb_branch", "Ha Noi Branch"))
    applicant = data.get("applicant_name", "")
    if applicant:
        cell_name = t0.rows[1].cells[0]
        _replace_in_cell(cell_name, "Tên công ty/Name of the company:",
                         f"Tên công ty/Name of the company: {applicant}")


def _fill_lc_type(doc: Document, data: dict) -> None:
    """Table 1 Row 0: LC type checkboxes and issuance method."""
    t1 = doc.tables[1]
    cell = t1.rows[0].cells[0]
    lc_type = (data.get("lc_type") or "Irrevocable").lower()
    method = (data.get("issuance_method") or "SWIFT").upper()

    if "transferable" in lc_type:
        _select_checkbox_in_cell(cell, "Irrevocable")
        _select_checkbox_in_cell(cell, "Transferable")
    elif "confirmed" in lc_type:
        _select_checkbox_in_cell(cell, "Irrevocable")
        _select_checkbox_in_cell(cell, "Confirmed")
    else:
        _select_checkbox_in_cell(cell, "Irrevocable")

    if method == "SWIFT":
        _select_checkbox_in_cell(cell, "Telex/SWIFT")
    else:
        _select_checkbox_in_cell(cell, "Mail")


def _fill_dates(doc: Document, data: dict) -> None:
    """Table 1 Rows 1-2: Expiry date, expiry place, latest shipment date."""
    t1 = doc.tables[1]
    if data.get("expiry_date"):
        _replace_in_cell(t1.rows[1].cells[0], "--/--/--", data["expiry_date"])
    if data.get("expiry_place"):
        cell = t1.rows[1].cells[0]
        for para in cell.paragraphs:
            ft = _full_text(para)
            if "Expiry Date" in ft:
                _set_para_text(para, ft + f"  Place: {data['expiry_place']}")
                break
    if data.get("latest_shipment_date"):
        _replace_in_cell(t1.rows[1].cells[2], "--/--/--", data["latest_shipment_date"])


def _fill_beneficiary_bank(doc: Document, data: dict) -> None:
    """Table 1 Row 2: Beneficiary's bank — add name/address as new paragraphs."""
    t1 = doc.tables[1]
    row = t1.rows[2]
    bank_name = data.get("beneficiary_bank_name", "")
    bank_addr = data.get("beneficiary_bank_address", "")
    bic = data.get("beneficiary_bank_bic", "")
    if bank_name or bank_addr:
        cell0 = row.cells[0]
        if bank_name:
            cell0.add_paragraph(bank_name)
        if bank_addr:
            cell0.add_paragraph(bank_addr)
    if bic:
        cell2 = row.cells[2]
        replaced = _replace_in_cell(cell2, "BIC code (preferably)", f"BIC/SWIFT: {bic}")
        if not replaced:
            cell2.add_paragraph(f"BIC/SWIFT: {bic}")


def _fill_applicant(doc: Document, data: dict) -> None:
    """Table 1 Row 3: Applicant full name and address."""
    t1 = doc.tables[1]
    row = t1.rows[3]
    name = data.get("applicant_name", "")
    addr = data.get("applicant_address", "")
    cell0 = row.cells[0]
    if name:
        cell0.add_paragraph(f"Full name: {name}")
    if addr:
        cell0.add_paragraph(f"Address: {addr}")


def _fill_beneficiary(doc: Document, data: dict) -> None:
    """Table 1 Row 4: Beneficiary full name, address, account."""
    t1 = doc.tables[1]
    row = t1.rows[4]
    name = data.get("beneficiary_name", "")
    addr = data.get("beneficiary_address", "")
    acct = data.get("beneficiary_account_no", "")
    cell0 = row.cells[0]
    if name:
        cell0.add_paragraph(f"Full name: {name}")
    if addr:
        cell0.add_paragraph(f"Address: {addr}")
    if acct:
        replaced = _replace_in_cell(row.cells[2], "Account No.", f"Account No. {acct}")
        if not replaced:
            row.cells[2].add_paragraph(f"Account No.: {acct}")


def _fill_amount(doc: Document, data: dict) -> None:
    """Table 1 Rows 5-6: Currency, amount, tolerance, amount in words."""
    t1 = doc.tables[1]
    row5 = t1.rows[5]
    currency = data.get("currency", "")
    amount = data.get("amount", "")
    tolerance = data.get("amount_tolerance", "0")
    if currency:
        _replace_in_cell(row5.cells[0], "Currency (ISO)", f"Currency (ISO): {currency}")
    if amount:
        _replace_in_cell(row5.cells[1], "Amount", f"Amount: {amount}")
    if tolerance:
        _replace_in_cell(row5.cells[2], "% More or Less Allowed",
                         f"% More or Less Allowed: {tolerance}%")
    words = data.get("amount_in_words", "")
    if words:
        _replace_in_cell(t1.rows[6].cells[0], "in words:", f"in words: {words}")


def _fill_draft(doc: Document, data: dict) -> None:
    """Table 1 Row 7: Draft type."""
    t1 = doc.tables[1]
    cell = t1.rows[7].cells[0]
    draft = (data.get("draft_type") or "Sight").lower()
    if "sight" in draft:
        _select_checkbox_in_cell(cell, "Sight")
    elif "usance" in draft or "days" in draft:
        days = data.get("draft_days", "")
        if days:
            _replace_in_para(cell.paragraphs[0], "days after Bill of Lading Date",
                             f"{days} days after Bill of Lading Date")
    elif "not required" in draft or "no draft" in draft:
        _select_checkbox_in_cell(cell, "Drafts not required")


def _fill_shipment_options(doc: Document, data: dict) -> None:
    """Table 1 Row 8: Partial shipment and transhipment.

    Partial shipment (cell0 para0): single paragraph, runs match exactly.
    Transhipment (cell2 para1): 'Not allowed' is split ('Not'/' '/'allowed') —
    use nth-checkbox positional selection instead.
    """
    t1 = doc.tables[1]
    row = t1.rows[8]
    partial = (data.get("partial_shipment") or "Not allowed").lower()
    trans = (data.get("transhipment") or "Not allowed").lower()

    cell0 = row.cells[0]
    if "not" in partial:
        _select_checkbox_in_cell(cell0, "Not allowed")
    else:
        _select_checkbox_in_cell(cell0, "Allowed")

    # cell2: checkbox[0]=Allowed, checkbox[1]=Not allowed
    cell2 = row.cells[2]
    if "not" in trans:
        _select_nth_checkbox_in_cell(cell2, 1)
    else:
        _select_nth_checkbox_in_cell(cell2, 0)


def _fill_ports(doc: Document, data: dict) -> None:
    """Table 1 Row 9: Ports."""
    t1 = doc.tables[1]
    cell = t1.rows[9].cells[0]
    loading = data.get("port_of_loading", "")
    discharge = data.get("port_of_discharge", "")
    if loading or discharge:
        port_text = (
            "(10) Shipment\n"
            f"Port of loading: {loading}\n"
            f"Port of discharge: {discharge}"
        )
        for para in cell.paragraphs:
            if "Shipment" in _full_text(para):
                _set_para_text(para, port_text)
                break


def _fill_incoterms(doc: Document, data: dict) -> None:
    """Table 1 Row 10: Incoterms."""
    t1 = doc.tables[1]
    cell = t1.rows[10].cells[0]
    inco = (data.get("incoterms") or "").upper()
    version = data.get("incoterms_version") or "2020"
    named_port = data.get("named_port") or data.get("port_of_discharge", "")
    if inco:
        _select_checkbox_in_cell(cell, inco)
    for para in cell.paragraphs:
        ft = _full_text(para)
        if "INCOTERMS" in ft.upper() or "Shipping Terms" in ft:
            new_text = ft.rstrip()
            if inco:
                new_text += f"\nSelected: {inco} {named_port} (INCOTERMS {version})"
            _set_para_text(para, new_text)
            break


def _fill_goods(doc: Document, data: dict) -> None:
    """Table 1 Row 11: Description of goods."""
    t1 = doc.tables[1]
    cell = t1.rows[11].cells[0]
    goods = data.get("description_of_goods", "")
    if goods:
        for para in cell.paragraphs:
            if "Description of goods" in _full_text(para):
                _set_para_text(para, f"(12) Description of goods and/or Services\n{goods}")
                break


def _fill_documents(doc: Document, data: dict) -> None:
    """Table 1 Row 12: Tick Wingdings checkboxes for required documents.

    Template paragraphs in T1R12 cell0:
      para0 = "Document required" (no checkbox)
      para1 = "This documentary credit is available against..." (no checkbox)
      para2 = \\uf06f signed commercial invoice
      para3 = \\uf06f full set ocean B/L
      para4 = \\uf06f air waybill
      para5 = \\uf06f Inspection certificate
      para6 = \\uf06f Certificate of quality
      para7 = \\uf06f full set insurance certificate/policy
      para8 = \\uf06f certificate of origin
      para9 = \\uf06f packing list
      para10 = \\uf06f Beneficiary's Certificate
      para11 = \\uf06f Other documents
    """
    t1 = doc.tables[1]
    cell = t1.rows[12].cells[0]
    docs_data = data.get("documents") or {}
    paras = cell.paragraphs

    # Map para index → data key (None means skip auto-tick)
    doc_para_map = {
        2: bool(docs_data.get("commercial_invoice")),
        3: bool(docs_data.get("bill_of_lading")),
        4: False,   # air waybill — only tick if explicitly present in other_documents
        5: bool(docs_data.get("inspection_certificate")),
        6: False,   # quality certificate — not a standard field
        7: bool(docs_data.get("insurance_certificate")),
        8: bool(docs_data.get("certificate_of_origin")),
        9: bool(docs_data.get("packing_list")),
        10: False,  # beneficiary cert — not a standard extracted field
        11: False,  # other documents — handled below
    }

    for para_idx, should_tick in doc_para_map.items():
        if should_tick and para_idx < len(paras):
            _check_run0_in_para(paras[para_idx])

    other_docs = docs_data.get("other_documents") or []
    if other_docs and len(paras) > 11:
        _check_run0_in_para(paras[11])
        # Append the list of other documents after the run text
        last_run = paras[11].runs[-1] if paras[11].runs else None
        if last_run and other_docs:
            last_run.text = last_run.text.rstrip() + ": " + "; ".join(other_docs)


def _fill_additional_conditions(doc: Document, data: dict) -> None:
    """Table 1 Row 13: Additional conditions and standard checkboxes.

    Template para structure:
      para0 = "Additional conditions:" (no checkbox — label only)
      para1 = \\uf06f Documents must be issued in English  (always tick)
      para2 = \\uf06f The amount utilized must be endorsed on the reverse (always tick)
    """
    t1 = doc.tables[1]
    cell = t1.rows[13].cells[0]
    paras = cell.paragraphs

    # Tick both standard checkboxes unconditionally
    if len(paras) > 1:
        _check_run0_in_para(paras[1])
    if len(paras) > 2:
        _check_run0_in_para(paras[2])

    # Append any custom conditions to para0 label
    conditions = (data.get("additional_conditions") or "").strip()
    if conditions and paras:
        # Filter out standard conditions already present in the template
        custom_lines = []
        for line in conditions.splitlines():
            line = line.strip()
            if line and "English" not in line and "amount utilized" not in line.lower():
                custom_lines.append(line)
        if custom_lines:
            full = _full_text(paras[0])
            _set_para_text(paras[0], full.rstrip() + "\n" + "\n".join(custom_lines))


def _fill_charges(doc: Document, data: dict) -> None:
    """Table 1 Row 14: Tick charge-responsibility checkboxes.

    Template para2 layout (4 Wingdings checkboxes):
      run0  = \\uf06f  → issuing bank / Applicant
      run4  = \\uf06f  → issuing bank / Beneficiary
      run8  = \\uf06f  → other banks  / Applicant
      run12 = \\uf06f  → other banks  / Beneficiary
    """
    t1 = doc.tables[1]
    cell = t1.rows[14].cells[0]
    issuing_for = (data.get("issuing_bank_charges_for") or "Applicant").lower()
    other_for = (data.get("other_bank_charges_for") or "Beneficiary").lower()

    paras = cell.paragraphs
    if len(paras) < 3:
        return
    runs = paras[2].runs

    def _tick(idx: int) -> None:
        if idx < len(runs) and runs[idx].text in _UNCHECKED:
            runs[idx].text = _CHECKED

    if "applicant" in issuing_for:
        _tick(0)   # issuing bank → Applicant
    else:
        _tick(4)   # issuing bank → Beneficiary

    if "applicant" in other_for:
        _tick(8)   # other banks → Applicant
    else:
        _tick(12)  # other banks → Beneficiary


def _fill_presentation_period(doc: Document, data: dict) -> None:
    """Table 2: Presentation period."""
    period = str(data.get("presentation_period") or "21")
    period_num = period.split()[0] if period else "21"
    t2 = doc.tables[2]
    cell = t2.rows[0].cells[0]
    if period_num == "21":
        # "21 days after shipment date" is fragmented across many runs —
        # replace Run 0 directly rather than relying on text-match search.
        for para in cell.paragraphs:
            if para.runs and any(c in para.runs[0].text for c in _UNCHECKED):
                for c in _UNCHECKED:
                    if c in para.runs[0].text:
                        para.runs[0].text = para.runs[0].text.replace(c, _CHECKED, 1)
                break
    else:
        _replace_in_cell(t2.rows[0].cells[1], "Other:", f"Other: {period_num} days after shipment date")


def _fill_upon_receipt(doc: Document, data: dict) -> None:
    """Table 3 Row 0: 'Upon receipt of' checkboxes.

    Template para structure in T3R0 cell0:
      para0 = 'Upon receipt of \\uf06f the Tested Telex/Swift'  (run6 = \\uf06f)
      para1 = '\\uf06f the Documents'                            (run0 = \\uf06f)

    SWIFT issuance → tick para0; always tick para1 (documents).
    """
    if len(doc.tables) <= 3:
        return
    t3 = doc.tables[3]
    if not t3.rows:
        return
    cell = t3.rows[0].cells[0]
    paras = cell.paragraphs
    method = (data.get("issuance_method") or "SWIFT").upper()

    # para0: "Upon receipt of  the Tested Telex/Swift"
    if paras and ("SWIFT" in method or "TELEX" in method):
        for run in paras[0].runs:
            if run.text in _UNCHECKED:
                run.text = _CHECKED
                break

    # para1: " the Documents" — always tick
    if len(paras) > 1:
        _check_run0_in_para(paras[1])


def _fill_fees_table(doc: Document, data: dict) -> None:
    """Table 6: Fee-responsibility checkboxes.

    Template layout (T6):
      row0 = headers: Domestic fees | Overseas fees | Confirmation fees
      row1 = Applicant  row: cell1=Domestic, cell2=Overseas, cell3=Confirmation
      row2 = Beneficiary row: cell1=Domestic, cell2=Overseas, cell3=Confirmation

    Each cell has para0 with run0=\\uf06f.  Tick based on charges_for fields:
      issuing_bank_charges_for → Domestic column
      other_bank_charges_for   → Overseas + Confirmation columns
    """
    if len(doc.tables) <= 6:
        return
    t6 = doc.tables[6]
    if len(t6.rows) < 3:
        return
    issuing_for = (data.get("issuing_bank_charges_for") or "Applicant").lower()
    other_for = (data.get("other_bank_charges_for") or "Beneficiary").lower()

    row_app = t6.rows[1]  # Applicant row
    row_ben = t6.rows[2]  # Beneficiary row

    def _tick_cell(row, col_idx: int) -> None:
        cells = row.cells
        if col_idx < len(cells) and cells[col_idx].paragraphs:
            _check_run0_in_para(cells[col_idx].paragraphs[0])

    # Domestic fees → issuing bank charges
    if "applicant" in issuing_for:
        _tick_cell(row_app, 1)
    else:
        _tick_cell(row_ben, 1)

    # Overseas fees → other banks' charges
    if "applicant" in other_for:
        _tick_cell(row_app, 2)
        _tick_cell(row_app, 3)  # Confirmation same as overseas
    else:
        _tick_cell(row_ben, 2)
        _tick_cell(row_ben, 3)


def _fill_contract_reference(doc: Document, data: dict) -> None:
    """Commitment paragraph (P33/P34): fill contract number and date.

    P33: 'số ..........…… ngày'  — Vietnamese side
    P34: 'Sales Contract No….. dated……'  — English side
    """
    contract_no = (data.get("contract_number") or "").strip()
    contract_date = (data.get("contract_date") or "").strip()

    if not contract_no and not contract_date:
        return

    # P33: replace placeholder run then append date to 'ngày' run
    for para in doc.paragraphs:
        full = _full_text(para)
        if '..........……' in full and 'ngày' in full:
            for run in para.runs:
                if '..........……' in run.text and contract_no:
                    run.text = run.text.replace('..........……', contract_no)
                elif run.text.strip() == 'ngày' and contract_date:
                    run.text = f'ngày {contract_date}'
            break

    # P34: 'Sales Contract No….. dated……'
    for para in doc.paragraphs:
        full = _full_text(para)
        if 'Sales Contract No' in full:
            for run in para.runs:
                if 'No…..' in run.text and contract_no:
                    run.text = run.text.replace('No…..', f'No. {contract_no}')
                if run.text.startswith('dated……') and contract_date:
                    run.text = run.text.replace('dated……', f'dated {contract_date} ', 1)
            break


def _replace_buyer_seller(doc: Document, data: dict) -> None:
    """Replace 'buyer'→'the applicant' and 'seller'→'the beneficiary' globally.

    Skips T1R11 (goods description) per user requirement — 'Remark' or goods
    description section should preserve original buyer/seller terminology.
    Per-run replacement preserves Wingdings checkbox characters in run 0.
    """
    t1 = doc.tables[1]
    goods_cell_id = id(t1.rows[11].cells[0])

    def _replace_runs(para) -> None:
        for run in para.runs:
            if not run.text:
                continue
            new = re.sub(r'\bbuyer\b', 'the applicant', run.text, flags=re.IGNORECASE)
            new = re.sub(r'\bseller\b', 'the beneficiary', new, flags=re.IGNORECASE)
            if new != run.text:
                run.text = new

    for para in doc.paragraphs:
        _replace_runs(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell) == goods_cell_id:
                    continue
                for para in cell.paragraphs:
                    _replace_runs(para)


# ─── Main public function ───────────────────────────────────────────────────

def fill_lc_template(
    data: dict,
    template_path: str,
    output_path: str,
) -> str:
    """Fill the Vietcombank LC application DOCX template.

    Args:
        data: LCApplicationData.model_dump() dict.
        template_path: Path to the blank DOCX template.
        output_path: Where to save the filled DOCX.

    Returns:
        output_path (str) on success.
    """
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    _fill_header(doc, data)
    _fill_lc_type(doc, data)
    _fill_dates(doc, data)
    _fill_beneficiary_bank(doc, data)
    _fill_applicant(doc, data)
    _fill_beneficiary(doc, data)
    _fill_amount(doc, data)
    _fill_draft(doc, data)
    _fill_shipment_options(doc, data)
    _fill_ports(doc, data)
    _fill_incoterms(doc, data)
    _fill_goods(doc, data)
    _fill_documents(doc, data)
    _fill_additional_conditions(doc, data)
    _fill_charges(doc, data)
    _fill_presentation_period(doc, data)
    _fill_upon_receipt(doc, data)
    _fill_fees_table(doc, data)
    _fill_contract_reference(doc, data)
    _replace_buyer_seller(doc, data)

    doc.save(output_path)
    return output_path
